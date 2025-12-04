import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches
from io import BytesIO, StringIO

# Set page configuration at the very top
st.set_page_config(page_title="Soil Mechanics Virtual Lab", layout="wide")

# Import ALL test modules, including ucs_test
from tabs import (
    sieve_analysis,
    liquid_limit_casagrande,
    liquid_limit_cone,
    plastic_limit,
    core_cutter,
    specific_gravity,
    constant_head,
    variable_head,
    light_compaction,
    direct_shear,
    ucs_test,  # UCS test now included
    consolidation,
    cbr_test,
    vane_shear,
    triaxial_test
)

# Available tests - UCS test added
test_functions = {
    "Sieve Analysis": sieve_analysis,
    "Liquid Limit (Casagrande)": liquid_limit_casagrande,
    "Liquid Limit (Cone Penetrometer)": liquid_limit_cone,
    "Plastic Limit": plastic_limit,
    "Core Cutter Method": core_cutter,
    "Specific Gravity": specific_gravity,
    "Constant Head Permeability": constant_head,
    "Variable Head Permeability": variable_head,
    "Light Compaction Test": light_compaction,
    "Direct Shear Test": direct_shear,
    "Unconfined Compressive Strength (UCS)": ucs_test,
    "Consolidation Test": consolidation,
    "California Bearing Ratio (CBR) Test": cbr_test,
    "Vane Shear Test": vane_shear,
    "Undrained Triaxial Test": triaxial_test
}

# Initialize session state
if "started" not in st.session_state:
    st.session_state.started = False
if "completed_tests" not in st.session_state:
    st.session_state.completed_tests = {}
if "current_result" not in st.session_state:
    st.session_state.current_result = None
if "last_test" not in st.session_state:
    st.session_state.last_test = None

# App Header
st.markdown("<h2 style='text-align: center;'>🧱 Soil Mechanics Virtual Laboratory</h2>", unsafe_allow_html=True)

# Start Button
if not st.session_state.started:
    if st.button("▶️ Start Soil Testing"):
        st.session_state.started = True
    st.stop()

# Test Selection
test_selection = st.selectbox("🔍 Select the test you want to perform:", list(test_functions.keys()))

# Run test and store result
returned_result = test_functions[test_selection].run()
if returned_result is not None:
    st.session_state.completed_tests[test_selection] = returned_result
    st.session_state.last_test = test_selection

# --- GLOBAL REPORT GENERATION (THIS SECTION WILL INCLUDE ALL COMPLETED TESTS) ---
if st.session_state.completed_tests:
    st.markdown("---")
    st.markdown("### 📄 Generate Test Report (All Completed Tests)")
    report_format = st.radio("Select report format", ("Excel", "Word (DOCX)"), key="global_report_format")

    if st.button("📥 Generate Combined Report", key="generate_global_report_button"):
        if report_format == "Excel":
            output = BytesIO()
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                for test_name, data in st.session_state.completed_tests.items():
                    if isinstance(data, dict):
                        for key, value in data.items():
                            if isinstance(value, pd.DataFrame):
                                sheet_name = f"{test_name[:20]}_{key[:10]}"
                                sheet_name = "".join(c for c in sheet_name if c.isalnum() or c in ['_', '-']).replace(' ', '_')
                                value.to_excel(writer, sheet_name=sheet_name, index=False)
                    elif isinstance(data, pd.DataFrame):
                        sheet_name = test_name[:31]
                        sheet_name = "".join(c for c in sheet_name if c.isalnum() or c in ['_', '-']).replace(' ', '_')
                        data.to_excel(writer, sheet_name=sheet_name, index=False)
            st.download_button(
                "📥 Download Excel Report (All Tests)",
                data=output.getvalue(),
                file_name="Soil_Test_Report_Combined.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

        elif report_format == "Word (DOCX)":
            doc = Document()
            doc.add_heading("🧱 Soil Testing Lab Report - Combined", level=0)

            for test_name, test_results_dict in st.session_state.completed_tests.items():
                doc.add_page_break()
                doc.add_heading(f"Test: {test_name}", level=1)

                if isinstance(test_results_dict, dict):
                    for section_name, section_data in test_results_dict.items():
                        if isinstance(section_data, pd.DataFrame):
                            doc.add_heading(f"{section_name}", level=2)
                            table = doc.add_table(rows=1, cols=len(section_data.columns))
                            table.style = 'Table Grid'
                            hdr_cells = table.rows[0].cells
                            for i, col in enumerate(section_data.columns):
                                hdr_cells[i].text = str(col)
                            for _, row in section_data.iterrows():
                                row_cells = table.add_row().cells
                                for i, val in enumerate(row):
                                    row_cells[i].text = str(val)
                            doc.add_paragraph("")

                        elif isinstance(section_data, BytesIO):
                            doc.add_heading(f"{section_name}", level=2)
                            try:
                                section_data.seek(0)
                                doc.add_picture(section_data, width=Inches(5))
                                doc.add_paragraph("")
                            except Exception as e:
                                doc.add_paragraph(f"Could not load image for {section_name}: {e}")
                        else:
                            doc.add_paragraph(f"**{section_name}:** {section_data}")
                elif isinstance(test_results_dict, pd.DataFrame):
                    doc.add_heading("Results Table", level=2)
                    table = doc.add_table(rows=1, cols=len(test_results_dict.columns))
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for i, col in enumerate(test_results_dict.columns):
                        hdr_cells[i].text = str(col)
                    for _, row in test_results_dict.iterrows():
                        row_cells = table.add_row().cells
                        for i, val in enumerate(row):
                            row_cells[i].text = str(val)
                    doc.add_paragraph("")

            buffer = BytesIO()
            doc.save(buffer)
            st.download_button(
                "📥 Download Word Report (All Tests)",
                data=buffer.getvalue(),
                file_name="Soil_Test_Report_Combined.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document

def run():
    st.subheader("🧪 Consolidation Test (IS 2720 Part 15:1986)")

    # Sample dimensions
    h0 = st.number_input("Initial height of soil sample (cm)", value=2.0, min_value=0.1)
    d = st.number_input("Diameter of soil sample (cm)", value=6.0, min_value=0.1)
    dial_lc = st.number_input("Dial gauge least count (mm/div)", value=0.01)

    A = (np.pi / 4) * d ** 2
    st.info(f"Cross-sectional Area: {A:.2f} cm²")

    # Initial void ratio input
    e0 = st.number_input("Initial void ratio (e₀)", value=0.9, min_value=0.0)

    # Input number of readings
    num_readings = st.number_input("Number of Load Increments", min_value=1, max_value=20, value=6, step=1)

    loads, init_readings, final_readings = [], [], []

    st.markdown("### 📥 Enter Data for Each Load Increment")
    for i in range(num_readings):
        st.markdown(f"#### 🔹 Load Increment {i+1}")
        col1, col2, col3 = st.columns(3)
        with col1:
            load = st.number_input(f"Load (kg/cm²) {i+1}", key=f"load_{i}", format="%.3f")
        with col2:
            init = st.number_input(f"Initial Reading (div) {i+1}", key=f"init_{i}", format="%.3f")
        with col3:
            final = st.number_input(f"Final Reading (div) {i+1}", key=f"final_{i}", format="%.3f")
        loads.append(load)
        init_readings.append(init)
        final_readings.append(final)

    if st.button("📊 Calculate"):
        try:
            df = pd.DataFrame({
                "Load (kg/cm²)": loads,
                "Initial Reading (div)": init_readings,
                "Final Reading (div)": final_readings
            })

            df["Settlement (mm)"] = (df["Final Reading (div)"] - df["Initial Reading (div)"]) * dial_lc
            df["Compression (cm)"] = df["Settlement (mm)"] / 10
            df["Strain"] = df["Compression (cm)"] / h0
            df["Void Ratio"] = e0 - df["Strain"] * (1 + e0)
            df = df[df["Void Ratio"] > 0]
            df["log(Load)"] = np.log10(df["Load (kg/cm²)"])

            st.markdown("### 📋 Result Table")
            st.dataframe(df.style.format(precision=4))

            # Plot e-log P
            st.markdown("### 📈 e – log(P) Curve")
            fig, ax = plt.subplots()
            ax.plot(df["log(Load)"], df["Void Ratio"], marker='o')
            ax.set_xlabel("log(Load) [log(kg/cm²)]")
            ax.set_ylabel("Void Ratio (e)")
            ax.set_title("e – log(P) Curve")
            ax.invert_yaxis()
            ax.grid(True)
            st.pyplot(fig)

            # Calculate Compression Index (Cc)
            if len(df) >= 3:
                slope, intercept = np.polyfit(df["log(Load)"].iloc[-3:], df["Void Ratio"].iloc[-3:], 1)
                Cc = -slope
                st.success(f"**Compression Index (Cc)** ≈ {Cc:.4f}")
            else:
                st.warning("Not enough points to calculate Compression Index.")

            # Coefficient of Volume Compressibility (mv)
            if len(df) >= 2:
                delta_e = df["Void Ratio"].iloc[1] - df["Void Ratio"].iloc[0]
                delta_sigma = (df["Load (kg/cm²)"].iloc[1] - df["Load (kg/cm²)"].iloc[0]) * 0.98
                mv = delta_e / (1 + e0) / delta_sigma
                mv_cm2_per_kg = mv * 10000 / 0.98
                st.success(f"**Coefficient of Volume Compressibility (mv)** ≈ {mv_cm2_per_kg:.6f} cm²/kg")
            else:
                st.warning("Not enough data to calculate mv.")

            if st.button("📥 Download Word Report"):
                doc = Document()
                doc.add_heading("Consolidation Test Report", 0)
                doc.add_paragraph(f"Initial Height: {h0} cm")
                doc.add_paragraph(f"Diameter: {d} cm")
                doc.add_paragraph(f"Area: {A:.2f} cm²")
                doc.add_paragraph(f"Initial Void Ratio (e₀): {e0}")

                table = doc.add_table(rows=1, cols=6)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Load (kg/cm²)"
                hdr_cells[1].text = "Initial Reading"
                hdr_cells[2].text = "Final Reading"
                hdr_cells[3].text = "Settlement (mm)"
                hdr_cells[4].text = "Void Ratio"
                hdr_cells[5].text = "log(Load)"

                for _, row in df.iterrows():
                    cells = table.add_row().cells
                    cells[0].text = f"{row['Load (kg/cm²)']:.2f}"
                    cells[1].text = f"{row['Initial Reading (div)']:.2f}"
                    cells[2].text = f"{row['Final Reading (div)']:.2f}"
                    cells[3].text = f"{row['Settlement (mm)']:.3f}"
                    cells[4].text = f"{row['Void Ratio']:.4f}"
                    cells[5].text = f"{row['log(Load)']:.4f}"

                if 'Cc' in locals():
                    doc.add_paragraph(f"\nCompression Index (Cc): {Cc:.4f}")
                if 'mv_cm2_per_kg' in locals():
                    doc.add_paragraph(f"Coefficient of Volume Compressibility (mv): {mv_cm2_per_kg:.6f} cm²/kg")

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    label="📥 Download Consolidation Report",
                    data=buffer,
                    file_name="Consolidation_Test_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        except Exception as e:
            st.error(f"⚠️ Error: {e}")

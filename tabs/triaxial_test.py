import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document

def run():
    st.subheader("🧪 Undrained Triaxial Test (IS 2720 Part 11:1971)")

    st.markdown("### 🧾 Enter Sample Details")
    sample_area = st.number_input("Cross-sectional Area of Sample (cm²)", value=38.0, min_value=1.0)
    sample_length = st.number_input("Length of Sample (cm)", value=7.6, min_value=1.0)

    st.markdown("### 🔢 Enter Number of Trials")
    num_trials = st.number_input("Number of Trials", min_value=1, max_value=10, value=3, step=1)

    st.markdown("### ✍️ Enter Data for Each Trial")
    trial_data = []
    for i in range(num_trials):
        st.markdown(f"#### Trial {i+1}")
        col1, col2, col3 = st.columns(3)
        with col1:
            confining_pressure = st.number_input(f"Confining Pressure (kg/cm²) [{i+1}]", key=f"cp_{i}")
        with col2:
            deviator_stress = st.number_input(f"Deviator Stress at Failure (kg/cm²) [{i+1}]", key=f"ds_{i}")
        with col3:
            angle = st.number_input(f"Failure Angle (°) [{i+1}]", value=90.0, key=f"fa_{i}")  # optional
        trial_data.append((confining_pressure, deviator_stress, angle))

    if st.button("📊 Calculate Results"):
        df = pd.DataFrame(trial_data, columns=["Confining Pressure", "Deviator Stress", "Failure Angle"])
        df["Major Principal Stress (σ₁)"] = df["Confining Pressure"] + df["Deviator Stress"]
        df["Minor Principal Stress (σ₃)"] = df["Confining Pressure"]

        # Calculate Mohr circles and fit line for cohesion and phi
        sigma1 = df["Major Principal Stress (σ₁)"]
        sigma3 = df["Minor Principal Stress (σ₃)"]
        center = (sigma1 + sigma3) / 2
        radius = (sigma1 - sigma3) / 2

        # Mohr circle plot
        st.markdown("### 📈 Mohr’s Circles")
        fig, ax = plt.subplots()
        for c, r in zip(center, radius):
            circle = plt.Circle((c, 0), r, fill=False)
            ax.add_patch(circle)
            ax.plot(c + r, 0, 'ro')  # σ1
            ax.plot(c - r, 0, 'bo')  # σ3

        ax.set_xlabel("Normal Stress (σ) [kg/cm²]")
        ax.set_ylabel("Shear Stress (τ) [kg/cm²]")
        ax.set_title("Mohr’s Circles for Triaxial Test")
        ax.axis('equal')
        ax.grid(True)
        ax.set_xlim(left=0)
        st.pyplot(fig)

        # Linear fit for failure envelope
        normal_stress = radius
        shear_stress = np.sqrt(radius**2)
        fit = np.polyfit(center, shear_stress, 1)
        cohesion = fit[1]
        phi_rad = np.arctan(fit[0])
        phi_deg = np.degrees(phi_rad)

        st.success(f"**Cohesion (c)** ≈ {cohesion:.4f} kg/cm²")
        st.success(f"**Angle of Internal Friction (ϕ)** ≈ {phi_deg:.2f}°")

        st.markdown("### 📋 Result Table")
        st.dataframe(df.style.format(precision=3))

        # Word Report Generation
        if st.button("📥 Download Word Report"):
            doc = Document()
            doc.add_heading("Undrained Triaxial Test Report", 0)
            doc.add_paragraph(f"Sample Area: {sample_area} cm²")
            doc.add_paragraph(f"Sample Length: {sample_length} cm")

            table = doc.add_table(rows=1, cols=5)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Trial"
            hdr_cells[1].text = "σ₃ (kg/cm²)"
            hdr_cells[2].text = "Deviator Stress (kg/cm²)"
            hdr_cells[3].text = "σ₁ (kg/cm²)"
            hdr_cells[4].text = "Failure Angle (°)"

            for i, row in df.iterrows():
                cells = table.add_row().cells
                cells[0].text = str(i+1)
                cells[1].text = f"{row['Confining Pressure']:.2f}"
                cells[2].text = f"{row['Deviator Stress']:.2f}"
                cells[3].text = f"{row['Major Principal Stress (σ₁)']:.2f}"
                cells[4].text = f"{row['Failure Angle']:.2f}"

            doc.add_paragraph(f"\nCohesion (c): {cohesion:.4f} kg/cm²")
            doc.add_paragraph(f"Angle of Internal Friction (ϕ): {phi_deg:.2f}°")

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.download_button(
                label="📥 Download Triaxial Report",
                data=buffer,
                file_name="Triaxial_Test_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

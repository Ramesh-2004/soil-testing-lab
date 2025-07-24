import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
import math

def generate_word_report(df, mean_ucs):
    doc = Document()
    doc.add_heading('Unconfined Compressive Strength (UCS) Test Report', 0)
    doc.add_paragraph('IS Code: IS 2720 Part 10:1991')
    doc.add_paragraph(f"Average UCS: {round(mean_ucs, 2)} kPa")

    doc.add_heading('Results Table', level=1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)

    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, item in enumerate(row):
            row_cells[i].text = str(item)

    # Add plot image to the Word report
    import tempfile
    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmpfile:
        fig, ax = plt.subplots()
        ax.bar(df["Trial"], df["UCS (kPa)"], color="lightgreen")
        ax.set_xlabel("Trial")
        ax.set_ylabel("UCS (kPa)")
        ax.set_title("Unconfined Compressive Strength")
        plt.tight_layout()
        fig.savefig(tmpfile.name)
        doc.add_paragraph("\nUCS Chart:")
        doc.add_picture(tmpfile.name, width=Inches(5))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def run():
    st.subheader("🧪 Unconfined Compressive Strength (UCS) Test - IS 2720 Part 10:1991")

    st.markdown("Enter details for each trial as per the lab manual:")
    num_trials = st.number_input("Number of Trials", min_value=1, max_value=10, value=3, step=1)

    results = []

    for i in range(num_trials):
        st.markdown(f"### Trial {i+1}")

        d = st.number_input(f"Diameter of specimen (mm) [Trial {i+1}]", key=f"d_{i}")
        l = st.number_input(f"Length of specimen (mm) [Trial {i+1}]", key=f"l_{i}")
        ring_constant = st.number_input(f"Proving ring constant (N/div) [Trial {i+1}]", key=f"k_{i}")
        ring_reading = st.number_input(f"Proving ring reading at failure (div) [Trial {i+1}]", key=f"r_{i}")

        if d > 0 and ring_constant > 0 and ring_reading > 0:
            area = math.pi * (d / 2) ** 2
            failure_load = ring_constant * ring_reading  # in Newtons
            ucs = (failure_load / area) * 1000  # Convert to kPa

            results.append({
                "Trial": i+1,
                "Diameter (mm)": d,
                "Length (mm)": l,
                "Area (mm²)": round(area, 2),
                "Proving Ring Constant (N/div)": ring_constant,
                "Ring Reading (div)": ring_reading,
                "Failure Load (N)": round(failure_load, 2),
                "UCS (kPa)": round(ucs, 2)
            })

    if st.button("📊 Calculate UCS"):
        if results:
            df = pd.DataFrame(results)
            st.markdown("### 📋 Result Table")
            st.dataframe(df)

            mean_ucs = df["UCS (kPa)"].mean()
            st.success(f"🔹 Average UCS = {round(mean_ucs, 2)} kPa")

            st.markdown("### 📈 UCS per Trial")
            fig, ax = plt.subplots()
            ax.bar(df["Trial"], df["UCS (kPa)"], color="lightgreen")
            ax.set_xlabel("Trial")
            ax.set_ylabel("UCS (kPa)")
            ax.set_title("Unconfined Compressive Strength")
            st.pyplot(fig)

            buffer = generate_word_report(df, mean_ucs)
            st.download_button(
                label="📄 Download UCS Word Report",
                data=buffer,
                file_name="UCS_Test_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.warning("Please enter valid values for all required fields.")

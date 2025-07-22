import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
import math
import json

def run():
    st.subheader("🔗 Direct Shear Test (IS 2720 Part 13:1986)")
    st.markdown("""
    This test is performed to determine the **shear strength parameters (cohesion 'c' and
    angle of internal friction 'phi')** of a soil sample under drained conditions.
    """)

    box_dim = st.number_input("Side Length of Shear Box (mm)", min_value=1.0, value=60.0, format="%.2f", key="box_dim")
    area = (box_dim / 10) ** 2  # in cm²
    st.info(f"Calculated Shear Area: {area:.2f} cm²")

    proving_ring_const = st.number_input("Proving Ring Constant (kg/div)", value=1.0, format="%.2f", key="prc")
    dial_lc = st.number_input("Dial Gauge Least Count (mm/div)", value=0.01, format="%.3f", key="dlc")

    st.markdown("---")
    st.write("### Test Setup and Readings")

    num_trials = st.number_input("Number of Normal Stress Trials", min_value=1, max_value=5, value=3, step=1, key="num_trials")
    
    # Define the number of horizontal deformation readings ONCE for all trials
    n_readings_per_trial = st.number_input("Number of Readings (Horizontal Deformation & Proving Ring)", min_value=2, value=10, key="n_readings_per_trial")

    st.markdown("---")
    st.write("#### Global Horizontal Deformation Readings (Same for all trials)")

    # Initialize/adjust session state for global horizontal deformations
    if "horizontal_deformations_list" not in st.session_state:
        st.session_state.horizontal_deformations_list = [0.0] * n_readings_per_trial
    elif len(st.session_state.horizontal_deformations_list) != n_readings_per_trial:
        # Resize if n_readings_per_trial changes, preserving existing values
        old_list = st.session_state.horizontal_deformations_list
        st.session_state.horizontal_deformations_list = [0.0] * n_readings_per_trial
        for k in range(min(len(old_list), n_readings_per_trial)):
            st.session_state.horizontal_deformations_list[k] = old_list[k]


    # Input fields for global horizontal deformation
    for k in range(n_readings_per_trial):
        st.session_state.horizontal_deformations_list[k] = st.number_input(
            f"Horizontal Deformation Reading {k+1} (div)",
            value=float(st.session_state.horizontal_deformations_list[k]),
            key=f"h_def_global_{k}",
            format="%.1f" # Assuming divisions can be decimals
        )

    st.markdown("---")
    st.write("#### Trial Specific Inputs")

    # Initialize/adjust session state for trial-specific inputs
    if "trial_inputs_detailed" not in st.session_state:
        st.session_state.trial_inputs_detailed = {}
    
    # Ensure all trials exist and their proving ring reading lists are correctly sized
    for i in range(num_trials):
        trial_key = f"trial_{i+1}"
        if trial_key not in st.session_state.trial_inputs_detailed:
            st.session_state.trial_inputs_detailed[trial_key] = {
                "sigma_n": (i+1)*0.5, # Default normal stress
                "proving_ring_readings": [0.0] * n_readings_per_trial
            }
        elif len(st.session_state.trial_inputs_detailed[trial_key]["proving_ring_readings"]) != n_readings_per_trial:
            # Resize proving ring readings if n_readings_per_trial changes
            old_pr_list = st.session_state.trial_inputs_detailed[trial_key]["proving_ring_readings"]
            st.session_state.trial_inputs_detailed[trial_key]["proving_ring_readings"] = [0.0] * n_readings_per_trial
            for k in range(min(len(old_pr_list), n_readings_per_trial)):
                st.session_state.trial_inputs_detailed[trial_key]["proving_ring_readings"][k] = old_pr_list[k]

    # Remove extra trials if num_trials decreased
    keys_to_remove = [k for k in st.session_state.trial_inputs_detailed.keys() if int(k.split('_')[1]) > num_trials]
    for key in keys_to_remove:
        del st.session_state.trial_inputs_detailed[key]

    # Input fields for each trial's specific data
    for i in range(num_trials):
        trial_key = f"trial_{i+1}"
        
        with st.expander(f"Trial {i+1} Details (Normal Stress: {st.session_state.trial_inputs_detailed[trial_key]['sigma_n']:.2f} kg/cm²)", expanded=True):
            # Normal Stress Input
            st.session_state.trial_inputs_detailed[trial_key]["sigma_n"] = st.number_input(
                f"Normal Stress σₙ (kg/cm²) - Trial {i+1}",
                value=float(st.session_state.trial_inputs_detailed[trial_key]["sigma_n"]),
                key=f"norm_stress_{i+1}",
                format="%.2f"
            )

            st.markdown(f"**Proving Ring Readings for Trial {i+1}**")
            # Proving Ring Readings Input for current trial
            for k in range(n_readings_per_trial):
                st.session_state.trial_inputs_detailed[trial_key]["proving_ring_readings"][k] = st.number_input(
                    f"Proving Ring Reading {k+1} (div) - Trial {i+1}",
                    value=float(st.session_state.trial_inputs_detailed[trial_key]["proving_ring_readings"][k]),
                    key=f"pr_reading_{i+1}_{k+1}",
                    format="%.2f"
                )

    st.markdown("---")

    # --- Save Inputs Button ---
    if st.button("💾 Save All Inputs", key="save_all_inputs_button"):
        saved_data = {
            "box_dim": box_dim,
            "proving_ring_const": proving_ring_const,
            "dial_lc": dial_lc,
            "num_trials": num_trials,
            "n_readings_per_trial": n_readings_per_trial,
            "horizontal_deformations_list": st.session_state.horizontal_deformations_list,
            "trial_inputs_detailed": st.session_state.trial_inputs_detailed
        }
        saved = json.dumps(saved_data, indent=2)
        st.download_button("📥 Download Saved Inputs (JSON)", data=saved, file_name="direct_shear_inputs_individual.json")

    # --- Calculate Results Button ---
    if st.button("🧮 Calculate Results"):
        if not st.session_state.horizontal_deformations_list or \
           all(h == 0.0 for h in st.session_state.horizontal_deformations_list):
            st.error("Please enter valid (non-zero) Horizontal Deformation readings.")
            return

        normal_stresses = []
        shear_stresses = []
        all_dfs = [] # To store individual trial DataFrames for report generation

        for i in range(num_trials):
            trial_key = f"trial_{i+1}"
            trial_data = st.session_state.trial_inputs_detailed.get(trial_key, {})
            sigma_n = trial_data.get("sigma_n", 0.0)
            proving_ring_readings = trial_data.get("proving_ring_readings", [])

            if not proving_ring_readings or all(pr == 0.0 for pr in proving_ring_readings):
                st.warning(f"Trial {i+1} has no proving ring readings. Skipping this trial for calculation.")
                continue # Skip if no proving ring data for this trial

            # Construct the DataFrame for this trial using global deformations and trial-specific PR readings
            df = pd.DataFrame({
                "Horizontal Deformation (div)": st.session_state.horizontal_deformations_list,
                "Proving Ring Reading (div)": proving_ring_readings
            })

            # Ensure all relevant columns have numeric values for calculation
            if not all(pd.to_numeric(df["Proving Ring Reading (div)"], errors='coerce').notna()) or \
               not all(pd.to_numeric(df["Horizontal Deformation (div)"], errors='coerce').notna()):
                st.error(f"Trial {i+1}: Please ensure all Proving Ring and Horizontal Deformation readings are valid numbers. Skipping this trial.")
                continue # Skip calculation for this trial if inputs are invalid

            df["Shear Force (kg)"] = df["Proving Ring Reading (div)"] * proving_ring_const
            df["Shear Stress (kg/cm²)"] = df["Shear Force (kg)"] / area
            df["Deformation (mm)"] = df["Horizontal Deformation (div)"] * dial_lc

            st.markdown(f"### 📊 Results for Trial {i+1} (Normal Stress: {sigma_n:.2f} kg/cm²)")
            st.dataframe(df) # Display the calculated dataframe for review

            fig, ax = plt.subplots(figsize=(8, 5))
            ax.plot(df["Deformation (mm)"], df["Shear Stress (kg/cm²)"], marker='o')
            ax.set_title(f"Trial {i+1}: Shear Stress vs Deformation")
            ax.set_xlabel("Deformation (mm)")
            ax.set_ylabel("Shear Stress (kg/cm²)")
            ax.grid(True)
            st.pyplot(fig)
            plt.close(fig) # Close the figure to free memory

            tau_max = df["Shear Stress (kg/cm²)"].max()
            shear_stresses.append(tau_max)
            normal_stresses.append(sigma_n)
            all_dfs.append(df.copy()) # Store the fully calculated DF for report

        if len(normal_stresses) >= 2: # Ensure at least two valid points for regression
            # Mohr-Coulomb Line Fit
            coeffs = np.polyfit(normal_stresses, shear_stresses, 1)
            phi_rad = math.atan(coeffs[0])
            phi_deg = math.degrees(phi_rad)
            cohesion = coeffs[1]

            st.markdown("### ✅ Final Results (After All Trials)")
            st.write(f"Cohesion (c): **{cohesion:.3f} kg/cm²**")
            st.write(f"Angle of Internal Friction (ϕ): **{phi_deg:.2f}°**")

            fig2, ax2 = plt.subplots(figsize=(8, 5))
            ax2.scatter(normal_stresses, shear_stresses, color='blue', label='Data Points')
            x_line = np.linspace(0, max(normal_stresses) * 1.1, 100)
            y_line = coeffs[0] * x_line + coeffs[1]
            ax2.plot(x_line, y_line, color='red', label=f'Mohr-Coulomb Fit (c={cohesion:.3f}, φ={phi_deg:.2f}°)')
            ax2.set_xlabel("Normal Stress σₙ (kg/cm²)")
            ax2.set_ylabel("Shear Stress τ (kg/cm²)")
            ax2.set_title("Shear Stress vs Normal Stress")
            ax2.grid(True)
            ax2.legend()
            st.pyplot(fig2)
            plt.close(fig2) # Close the figure to free memory


            # Report Generation
            if st.button("📄 Generate Combined Word Report"):
                doc = Document()
                doc.add_heading("Direct Shear Test Report", 0)
                doc.add_paragraph(f"Shear Box Size: {box_dim} mm")
                doc.add_paragraph(f"Proving Ring Constant: {proving_ring_const} kg/div")
                doc.add_paragraph(f"Dial Gauge LC: {dial_lc} mm/div")
                doc.add_paragraph(f"Area: {area:.2f} cm²")
                doc.add_paragraph(f"**Cohesion (c): {cohesion:.3f} kg/cm²**")
                doc.add_paragraph(f"**Angle of Internal Friction (ϕ): {phi_deg:.2f}°**")
                
                # Add the main Shear Stress vs Normal Stress plot
                doc.add_heading("Overall Shear Strength Parameters Plot", level=1)
                buf_combined_plot = BytesIO()
                fig2.savefig(buf_combined_plot, format="png", dpi=300, bbox_inches="tight")
                buf_combined_plot.seek(0)
                doc.add_picture(buf_combined_plot, width=Inches(6))
                doc.add_page_break()


                for idx_trial in range(len(all_dfs)): # Iterate through collected valid trials
                    df_trial = all_dfs[idx_trial] # Use the stored calculated DF
                    doc.add_heading(f"Trial {idx_trial+1} (Normal Stress: {normal_stresses[idx_trial]:.2f} kg/cm²)", level=2)
                    
                    # Add data table
                    table = doc.add_table(rows=1, cols=len(df_trial.columns))
                    table.autofit = True
                    # Add headers dynamically
                    hdr_cells = table.rows[0].cells
                    for col_idx, col_name in enumerate(df_trial.columns):
                        hdr_cells[col_idx].text = col_name

                    # Add data rows
                    for row_idx in range(len(df_trial)):
                        row_cells = table.add_row().cells
                        for col_idx, col_name in enumerate(df_trial.columns):
                            value = df_trial[col_name].iloc[row_idx]
                            # Format numerical columns for better readability
                            if isinstance(value, (int, float)):
                                row_cells[col_idx].text = f"{value:.3f}"
                            else:
                                row_cells[col_idx].text = str(value)
                    doc.add_paragraph("\n") # Add a small space after table

                    # Add individual trial plot
                    buf_trial_plot = BytesIO()
                    temp_fig, temp_ax = plt.subplots(figsize=(8, 5))
                    temp_ax.plot(df_trial["Deformation (mm)"], df_trial["Shear Stress (kg/cm²)"], marker='o')
                    temp_ax.set_title(f"Trial {idx_trial+1}: Shear Stress vs Deformation (Normal Stress: {normal_stresses[idx_trial]:.2f} kg/cm²)")
                    temp_ax.set_xlabel("Deformation (mm)")
                    temp_ax.set_ylabel("Shear Stress (kg/cm²)")
                    temp_ax.grid(True)
                    temp_fig.savefig(buf_trial_plot, format="png", dpi=300, bbox_inches="tight")
                    plt.close(temp_fig) # Close temp figure
                    buf_trial_plot.seek(0)
                    doc.add_picture(buf_trial_plot, width=Inches(6))
                    doc.add_page_break() # Start new trial on a new page

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    "📥 Download Word Report",
                    data=buffer.getvalue(),
                    file_name="Direct_Shear_Test_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.warning("Please ensure at least two trials have valid data to calculate Cohesion and Angle of Internal Friction.")

    return None
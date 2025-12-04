import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches

def generate_cbr_report(data):
    doc = Document()
    doc.add_heading("California Bearing Ratio (CBR) Test Report", 0)

    doc.add_heading("General Observations", level=1)
    doc.add_paragraph(f"Mould Diameter: {data['Mould Diameter']} mm")
    doc.add_paragraph(f"Mould Height: {data['Mould Height']} mm")
    doc.add_paragraph(f"Proving Ring Constant: {data['Ring Constant']} kg/div")
    doc.add_paragraph(f"Initial Penetration Offset: {data['Offset']} mm")

    doc.add_heading("CBR Results", level=1)
    doc.add_paragraph(f"CBR at 2.5 mm: {data['CBR at 2.5 mm (%)']:.2f} %")
    doc.add_paragraph(f"CBR at 5.0 mm: {data['CBR at 5.0 mm (%)']:.2f} %")
    doc.add_paragraph(f"Final CBR: {data['Final CBR (%)']:.2f} %")
    doc.add_paragraph(f"Remark: {data['Remark']}")

    doc.add_heading("CBR Load vs. Penetration Data", level=1)
    table = doc.add_table(rows=1, cols=len(data['DataFrame'].columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(data['DataFrame'].columns):
        hdr_cells[i].text = col

    for index, row in data['DataFrame'].iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = f"{value:.2f}" if isinstance(value, float) else str(value)

    doc.add_page_break()
    doc.add_heading("Load vs. Penetration Curve", level=1)

    if data["Curve Plot"]:
        image_stream = data["Curve Plot"]
        doc.add_picture(image_stream, width=Inches(5.5))

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    output.name = "CBR_Test_Report.docx"
    return output

def run():
    st.title("California Bearing Ratio (CBR) Test")

    with st.expander("🔧 Enter Test Parameters"):
        mould_diameter = st.number_input("Mould Diameter (mm)", min_value=1.0, value=150.0)
        mould_height = st.number_input("Mould Height (mm)", min_value=1.0, value=127.0)
        ring_constant = st.number_input("Proving Ring Constant (kg/div)", min_value=0.01, value=0.5)
        offset = st.number_input("Initial Penetration Offset (mm)", min_value=0.0, value=0.5)

    st.markdown("---")
    st.subheader("🔢 Enter Load vs. Penetration Data")

    default_data = {
        "Penetration (mm)": [0.5, 1.0, 1.5, 2.0, 2.5, 3.0, 4.0, 5.0, 7.5, 10.0, 12.5],
        "Dial Gauge Reading (divisions)": [2, 4, 7, 10, 12, 14, 18, 22, 30, 38, 45]
    }

    df = pd.DataFrame(default_data)
    edited_df = st.data_editor(df, num_rows="dynamic")

    if st.button("📈 Calculate CBR"):
        edited_df["Load (kg)"] = edited_df["Dial Gauge Reading (divisions)"] * ring_constant
        edited_df["CBR (%)"] = (edited_df["Load (kg)"] / 1370) * 100

        cbr_2_5 = edited_df.loc[edited_df["Penetration (mm)"] == 2.5, "CBR (%)"].values[0] if 2.5 in edited_df["Penetration (mm)"].values else 0
        cbr_5_0 = edited_df.loc[edited_df["Penetration (mm)"] == 5.0, "CBR (%)"].values[0] if 5.0 in edited_df["Penetration (mm)"].values else 0
        final_cbr = max(cbr_2_5, cbr_5_0)
        remark = "Suitable for subgrade" if final_cbr > 10 else "Not suitable for subgrade"

        st.success(f"CBR at 2.5 mm: {cbr_2_5:.2f}%")
        st.success(f"CBR at 5.0 mm: {cbr_5_0:.2f}%")
        st.success(f"Final CBR (%): {final_cbr:.2f}%")
        st.info(f"Remark: {remark}")

        # Plot
        fig, ax = plt.subplots()
        ax.plot(edited_df["Penetration (mm)"], edited_df["Load (kg)"], marker='o')
        ax.set_xlabel("Penetration (mm)")
        ax.set_ylabel("Load (kg)")
        ax.set_title("Load vs Penetration Curve")
        ax.grid(True)

        image_stream = BytesIO()
        plt.savefig(image_stream, format='png')
        image_stream.seek(0)

        st.pyplot(fig)

        report_data = {
            "Mould Diameter": mould_diameter,
            "Mould Height": mould_height,
            "Ring Constant": ring_constant,
            "Offset": offset,
            "CBR at 2.5 mm (%)": cbr_2_5,
            "CBR at 5.0 mm (%)": cbr_5_0,
            "Final CBR (%)": final_cbr,
            "Remark": remark,
            "DataFrame": edited_df,
            "Curve Plot": image_stream
        }

        report_buffer = generate_cbr_report(report_data)

        st.download_button(
            "📥 Download Word Report",
            data=report_buffer,
            file_name="CBR_Test_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
